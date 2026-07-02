# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ---- Style setup ----
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Microsoft YaHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        if level == 0:
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        elif level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
        elif level == 3:
            run.font.size = Pt(11.5)
            run.font.color.rgb = RGBColor(0x34, 0x49, 0x5e)
    return h

def add_para(text, bold=False, italic=False, size=10.5, color=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(10.5)
    return p

def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F5F5F5')
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.font.name = 'Microsoft YaHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr_cells[i], '2C3E50')
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = ''
            p = row_cells[i].paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Microsoft YaHei'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.size = Pt(9.5)
    doc.add_paragraph()
    return table

def add_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'E8F4FD')
    p._p.get_or_add_pPr().append(shading)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), '2980B9')
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(10.5)
    run.italic = True
    run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
    return p

# ============================================================
# DOCUMENT CONTENT
# ============================================================

add_heading('NSEAP 知识认知细胞（Knowledge Cognitive Cell）', 0)
add_para('Elite20 Builder Program \u00b7 第 6 组 Knowledge Team', bold=True, size=11, color=(0x7f, 0x8c, 0x8d))
add_para('版本：v0.1 MVP \u00b7 对应能力等级：L2（标准化知识库），迈向 L3（Agent-Ready）', size=10, color=(0x7f, 0x8c, 0x8d), space_after=12)

# ---- Section 1 ----
add_heading('一、产品是什么', 1)
add_para('NSEAP AI Learning Operating System 共有六大模块：课程、挑战库、学习平台、Agent 库、本体体系、知识库。知识认知细胞（Knowledge Cognitive Cell）是第六模块——知识层。')
add_para('它不是普通文件夹，也不是文档目录，而是一个有 Identity、Capability、Interface 的结构化知识仓库。它的产品公式是：', space_after=4)

add_code_block(
    'Knowledge Cognitive Cell\n'
    '= Knowledge Repository          （知识仓库）\n'
    '+ Prompt Studio Seed            （提示词工作室雏形）\n'
    '+ Metadata Search               （元数据优先检索）\n'
    '+ Relationship Management       （知识关系管理）\n'
    '+ Agent Retrieval Context       （Agent 检索上下文）\n'
    '+ Knowledge Growth Loop         （知识增长闭环）'
)

add_heading('当前形态', 3)
add_para('v0.1 MVP——静态 Demo App + 结构化 Markdown 知识源 + JSON Schema + 模板 + 搜索索引 + 设计文档。验证的是知识模型、检索方式、关系结构和演示流程，不是最终产品形态。')

add_heading('成熟度定位', 3)
add_table(
    ['等级', '名称', '说明'],
    [
        ['L0', '散落文档', '文件存在，无共享结构'],
        ['L1', '文档库', '文件分目录，人可阅读'],
        ['L2', '标准化知识库', '模板 + metadata + 稳定分类 + GitHub Review ← 当前 MVP'],
        ['L3', 'Agent-Ready 知识库', 'Agent 可按 type/audience/tags/concepts/skills 检索 ← 下一目标'],
        ['L4', '知识认知细胞', '有明确 Identity / Capability / Interface / Contract'],
        ['L5', '进化型知识细胞', '使用反馈 + 评估 + 长期记忆 + KSTAR 自我改进 → Knowledge Librarian Agent'],
    ]
)

add_heading('信息架构', 3)
add_code_block(
    'Knowledge Repository\n'
    '\u251c\u2500 00-overview        系统说明与使用指南\n'
    '\u251c\u2500 01-course          课程知识（Syllabus / Weekly Plan / Lecture Notes）\n'
    '\u251c\u2500 02-challenges      挑战与任务（目标 / 步骤 / 交付物 / 评分 / 常见错误）\n'
    '\u251c\u2500 03-prompts         可复用提示词\n'
    '\u251c\u2500 04-faq             常见问题（学生 / 教师 / Builder）\n'
    '\u251c\u2500 05-best-practices  最佳实践\n'
    '\u251c\u2500 06-projects        项目案例与复盘\n'
    '\u251c\u2500 07-agents          Agent 上下文与 Knowledge Librarian\n'
    '\u2514\u2500 (未来) Rubric / Concept / Skill  评价标准与概念技能节点'
)

add_heading('核心数据模型（KnowledgeItem）', 3)
add_para('每一条知识都是一个结构化实体，不是一篇普通文档：')
add_table(
    ['字段类别', '字段', '作用'],
    [
        ['标识', 'id, title, type, status', '稳定 ID + 类型 + 状态（draft/review/stable/deprecated）'],
        ['检索', 'keywords, tags, concepts, skills, audience', 'metadata-first 搜索核心'],
        ['关系', 'related, relationships', '显式表达条目间关系（includes/requires/supports/usesPrompt/assessedBy/relatedTo）'],
        ['知识增长', 'situation, ontology, workflow, skill, evaluation, knowledgeGrowth', '串联"场景→本体→流程→技能→评估→学习→知识增长"闭环'],
        ['Agent 就绪', 'agentNotes', '告诉 Agent 何时检索、能回答什么问题、需要什么关联上下文、有什么限制'],
        ['审计', 'createdAt, updatedAt, createdBy, updatedBy', '版本与贡献追踪'],
    ]
)

# ---- Section 2 ----
add_heading('二、给谁用的（目标用户）', 1)
add_table(
    ['角色', '核心需求', '在产品里做什么'],
    [
        ['学习者', '查找课程说明、挑战任务、提示词模板、FAQ、项目案例', '在 Demo App 里按分类筛选 + 关键词搜索，找到可直接复用的知识'],
        ['教师 / 课程组织者', '维护课程知识，沉淀常见问题、优秀案例和教学经验', '用模板创建知识条目 → 填 metadata → 提交 GitHub → Review 后发布为 stable'],
        ['Builder / 项目成员', '按标准模板提交新知识条目，通过协作流程完成评审', '选知识类型 → 填写内容 → 校验 metadata → review → Merge 发布'],
        ['Agent', '根据问题、角色、课程、任务和上下文检索可用知识', '调用 Knowledge Repository API → 按 metadata + relationships 检索 → 生成带引用的回答'],
        ['后端 / 平台开发者', '按统一模型实现数据库、API、搜索、关系和 Agent 接口', '按 KnowledgeItem 模型建表 → 实现 metadata-first 搜索 → 预留 Agent 检索 API'],
    ]
)

# ---- Section 3 ----
add_heading('三、解决什么问题', 1)
add_table(
    ['#', '痛点', '现状', '解决方案'],
    [
        ['1', '资料分散', '课程说明、挑战任务、提示词、FAQ、项目案例分布在不同位置', '统一 Knowledge Repository，10 个分类目录，所有知识归一'],
        ['2', '复用困难', '好的提示词、案例和经验很难被后续学习者、教师、Builder 或 Agent 找到', 'metadata-first 搜索 + 稳定 ID + 关键词/标签/概念/技能多维检索'],
        ['3', '结构不足', '普通文档适合阅读，但不适合筛选、搜索、关联和系统调用', 'KnowledgeItem 数据模型 + JSON Schema 校验 + 模板标准化'],
        ['4', '关系缺失', '课程、挑战、技能、概念、提示词、项目之间的关系没有被显式表达', 'Relationship 表（source→predicate→target），6 种标准关系谓词'],
        ['5', '难以增长', '项目产出如果没有沉淀机制，很难变成后续可复用知识', '知识增长闭环 + KSTAR 更新流程'],
        ['6', 'Agent 难调用', '没有稳定 ID、metadata 和关系结构，Agent 很难可靠检索上下文', 'agentNotes 字段 + audience/type/status/related 显式化 + 未来 API'],
    ]
)
add_quote('一句话：解决"AI 教育项目中的知识无法被结构化存储、精准检索、显式关联、持续增长和被 Agent 可靠调用"的问题。')

# ---- Section 4 ----
add_heading('四、产品价值', 1)

add_heading('对 NSEAP 生态的价值', 2)
add_table(
    ['服务对象', '提供的价值'],
    [
        ['Curriculum Team（课程组）', '把 weekly plan、learning objectives、lecture notes 转成可复用课程知识条目'],
        ['Challenge Team（挑战组）', '把 Challenge 转成结构化条目，含目标、步骤、交付物、评分方式、常见错误、相关提示词和知识增长'],
        ['Agent Team（Agent 组）', '提供结构化上下文、提示词样例、FAQ 和 metadata，让 Agent 更干净地检索知识'],
        ['Ontology Team（本体组）', '提供 concepts、skills、relationships，作为未来 ontology node/edge 的输入'],
        ['Platform Team（平台组）', '提供干净的 Markdown 内容源和静态 Demo，未来可接 GitHub Pages、文档门户、LMS、API'],
        ['Demo Team（演示组）', '提供清楚的展示故事：项目知识不再散落，而是结构化、可检索、可复用的知识资产'],
    ]
)

add_heading('核心产品价值', 2)
add_table(
    ['价值点', '说明'],
    [
        ['不是文件夹，是认知细胞', '有 Identity / Capability / Interface，可演化为 Knowledge Librarian Agent'],
        ['metadata-first 搜索', '搜"挑战"优先返回 type=challenge 的条目，而非正文随机提到"挑战"的文档'],
        ['Agent-ready 设计', '通过 agentNotes + 显式 metadata + 关系结构，为未来 Agent 检索铺好地基'],
        ['可演化路径清晰', 'L0→L5 六级成熟度模型 + V0.1→V1.0 四阶段版本路线'],
        ['对齐 IEEE 标准体系', 'P2807.8（知识图谱）/ P3394（Agent Interface）/ P3428（能力等级）/ CognitiveCell / FDE Workbench'],
    ]
)

add_heading('版本路线', 2)
add_table(
    ['版本', '内容', '目标'],
    [
        ['V0.1（当前）', '静态 Demo + Markdown + JSON + 搜索索引 + 模板 + Schema + 设计文档', '验证产品模型、知识结构和演示流程'],
        ['V0.2', '更多知识条目 + 更完整 metadata + 自动索引生成 + metadata 校验', '让 Builder 可以持续补充知识'],
        ['V0.3', 'KnowledgeItem API + 数据库 + 搜索服务 + Markdown 导入导出 + 关系管理', '从静态 MVP 升级为可维护服务'],
        ['V1.0', '正式后端 + 工作台前端 + 审核流程 + Prompt Studio + Agent 检索接口 + FDE Workbench 集成', '成为 NSEAP 的正式知识仓库能力'],
    ]
)

# ---- Section 5 ----
add_heading('五、如何使用 — 快速上手', 1)

add_heading('方式一：跑起来看 Demo（任何人，2 分钟）', 2)
add_code_block(
    'cd nseap-knowledge-base/app\n'
    'python -m http.server 8000\n'
    '# 打开浏览器访问 http://localhost:8000'
)
add_para('Demo 功能：', bold=True, space_after=4)
add_bullet('浏览知识条目列表')
add_bullet('按 Course / Challenge / Prompt / Project / Agent 分类筛选')
add_bullet('搜索标题、标签、概念、技能、适用对象和 Markdown 文件内容')
add_bullet('查看条目的 metadata、concepts、skills、relationships')
add_bullet('展示知识增长流程：场景→本体→流程→技能→评估→学习→知识增长')
add_bullet('从 Demo 跳回源 Markdown 文件')

add_heading('方式二：作为 Builder 贡献知识条目（开发者）', 2)
add_code_block(
    'Step 1  选知识类型 → 找到对应目录（如 knowledge-base/02-challenges/）\n'
    'Step 2  复制模板  → templates/knowledge-item-template.md\n'
    'Step 3  填写内容  → 标题、摘要、keywords、tags、concepts、skills、\n'
    '                   related、relationships、situation→knowledgeGrowth、agentNotes\n'
    'Step 4  校验      → 对照 schemas/knowledge-item.schema.json 确认 metadata 合规\n'
    'Step 5  重新索引  → node scripts/build-search-index.js\n'
    'Step 6  提交      → GitHub PR → Peer Review → Agent Review → Merge → Release'
)

add_heading('方式三：作为后端开发者接续开发（平台组）', 2)
add_table(
    ['层面', '设计'],
    [
        ['数据库', 'knowledge_items + 5 张多值子表 + knowledge_relationships + knowledge_revisions'],
        ['API', 'GET/POST/PATCH/DELETE /api/knowledge-items + 关系接口 + 搜索建议 + Markdown 导入导出'],
        ['搜索', 'metadata-first：keywords > title > tags > concepts > skills > type > audience > headings > summary > content'],
        ['同步', 'V0.1 Markdown 为源 → V0.3 双向同步 → V1.0 数据库为主、Markdown 作为导出/审计格式'],
    ]
)
add_para('MVP 与后端对应关系：', bold=True, space_after=4)
add_table(
    ['当前 MVP', '正式后端'],
    [
        ['knowledge-base/*.md', 'knowledge_items.content_markdown + metadata 表'],
        ['app/knowledge-data.json', 'GET /api/knowledge-items'],
        ['app/search-index.json', '搜索索引表或搜索服务'],
        ['分类按钮', 'type 查询参数'],
        ['搜索框', '/api/knowledge-items?q=...'],
        ['关系展示', 'knowledge_relationships'],
        ['知识卡片模拟', 'POST /api/knowledge-items'],
    ]
)

# ---- Section 6 ----
add_heading('六、竞品分析', 1)

add_heading('竞品全景对比', 2)
add_table(
    ['维度', 'Confluence/Notion', 'Obsidian/Logseq', 'LangChain/Dify', 'Knowledge Cognitive Cell'],
    [
        ['定位', '团队协作文档库', '个人笔记网络', 'AI 知识检索层', '教育知识认知细胞'],
        ['结构化 metadata', '弱（自由格式）', '中（frontmatter）', '中（chunk+embedding）', '强（模型+Schema）'],
        ['知识关系', '无/弱（链接）', '中（双链）', '弱（向量相似度）', '强（显式+6种谓词）'],
        ['搜索方式', '全文搜索', '全文+链接', '向量语义搜索', 'metadata-first+全文'],
        ['Agent 就绪', '✗', '✗', '✓（知识无结构）', '✓（agentNotes+metadata）'],
        ['教育领域建模', '✗', '✗', '✗', '✓（原生类型）'],
        ['知识增长闭环', '✗', '✗', '✗', '✓'],
        ['演化路径', '无', '插件扩展', '无成熟度模型', 'L0-L5+V0.1-V1.0'],
        ['标准对齐', '无', '无', '无', 'IEEE P2807.8等'],
        ['部署方式', 'SaaS/自部署', '本地', '框架（自己搭）', '开源→后端→FDE'],
    ]
)

add_heading('竞品逐一分析', 2)

add_heading('1. 团队文档/Wiki（Confluence / Notion / GitHub Wiki）', 3)
add_table(
    ['维度', 'Confluence/Notion', 'Knowledge Cognitive Cell'],
    [
        ['定位', '通用团队协作文档库', '教育知识认知细胞'],
        ['知识结构', '自由格式页面，无强制 metadata', 'KnowledgeItem 模型 + JSON Schema 强制校验'],
        ['搜索', '全文搜索——搜"挑战"返回所有正文提到"挑战"的页面', 'metadata-first——搜"挑战"优先返回 type=challenge 的条目'],
        ['关系', '页面间链接，无语义关系', '显式 Relationship 表，6 种谓词'],
        ['Agent 就绪', '✗', '✓ agentNotes + 显式 metadata'],
        ['相同点', '都支持分类目录、协作、Review 流程', ''],
        ['关键差异', '给人看的文档库', '给人+Agent 共用的结构化知识仓库'],
    ]
)

add_heading('2. 个人知识管理（Obsidian / Logseq / Roam Research）', 3)
add_table(
    ['维度', 'Obsidian/Logseq', 'Knowledge Cognitive Cell'],
    [
        ['定位', '个人双链笔记网络', '组织级教育知识仓库'],
        ['关系', '双链（backlink）——隐式关联', '显式关系表——predicate 有语义'],
        ['metadata', 'frontmatter 可选，无强制 schema', 'KnowledgeItem JSON Schema 强制 + 模板'],
        ['搜索', '全文 + 标签', 'metadata-first 多维检索'],
        ['协作', '个人为主，Git 同步为辅', 'GitHub PR + Peer Review + Agent Review'],
        ['教育建模', '✗', '✓ 原生支持教育知识类型'],
        ['演化', '插件生态，无成熟度模型', 'L0-L5 认知细胞演化路径'],
        ['关键差异', '个人笔记网络', '组织级、Agent-ready、教育领域知识基础设施'],
    ]
)

add_heading('3. AI 知识库 / RAG 框架（LangChain / LlamaIndex / Dify KB）', 3)
add_table(
    ['维度', 'LangChain/Dify KB', 'Knowledge Cognitive Cell'],
    [
        ['定位', 'AI 应用的知识检索层', '教育知识认知细胞（人+Agent+平台共用）'],
        ['知识结构', '文档切块→向量化→语义检索，无类型/关系', 'KnowledgeItem 结构化实体 + 显式关系 + 知识增长闭环'],
        ['搜索', '向量语义搜索（语义相似）', 'metadata-first（意图精准匹配）——两者可互补'],
        ['Agent 就绪', '✓ 但知识是无结构的 chunk', '✓ 且知识是有结构的条目（agentNotes）'],
        ['教育领域', '✗ 通用', '✓ course/challenge/prompt/FAQ 原生类型'],
        ['知识增长', '✗ 无闭环机制', '✓ 场景→本体→...→知识增长 + KSTAR'],
        ['关键差异', '解决"Agent 怎么检索"', '解决"知识怎么被结构化组织、持续增长、演化成认知细胞"——在 RAG 上游'],
    ]
)

add_heading('4. 学习内容管理系统（LCMS / 传统教育内容管理平台）', 3)
add_table(
    ['维度', '传统 LCMS', 'Knowledge Cognitive Cell'],
    [
        ['定位', '教育内容管理（管课件、题库）', '教育知识认知细胞（管知识条目+关系+Agent 上下文）'],
        ['内容粒度', '课件/课程级别', '知识条目级别（可细到一条 Prompt、一条 FAQ）'],
        ['关系', '课程树（包含关系）', '知识图谱（6 种语义关系）'],
        ['Agent', '✗', '✓ agentNotes + 检索接口'],
        ['标准', 'SCORM 等传统教育标准', 'IEEE P2807.8/P3394/P3428 + Cognitive Cell'],
        ['关键差异', '管"内容分发"', '管"知识沉淀、关联、增长和 Agent 就绪"'],
    ]
)

# ---- Section 7 ----
add_heading('七、总结', 1)
add_quote(
    'NSEAP 知识认知细胞不是文件夹，不是文档库，而是一个有 Identity / Capability / Interface 的结构化知识仓库。'
    '它用 KnowledgeItem 模型 + metadata-first 搜索 + 显式关系 + 知识增长闭环 + agentNotes，'
    '把课程、挑战、提示词、FAQ、最佳实践和项目案例沉淀为可检索、可复用、可被 Agent 调用、'
    '未来可图谱化和可自我进化的认知资产。'
    '当前 v0.1 是 L2 标准化知识库的 MVP 验证，目标是演化到 L5 Knowledge Librarian Agent。'
)

# ---- Appendix ----
add_heading('附录：参考文件', 1)
add_table(
    ['文件', '对设计的影响'],
    [
        ['Elite20-Vibe-Coding-Course.docx', '明确 Builder Program 的协作、GitHub 提交、评审、文档化和发布流程'],
        ['CognitiveCell.docx', '将知识库设计为有身份、能力、接口和演化路径的认知细胞'],
        ['[Clean]P2807.8...docx', '为教育知识图谱、学习路径、本体和语义关系预留结构'],
        ['P3394-D1.0.0-IEEE-Draft', '为未来 Agent 接口、manifest、消息和互操作能力预留方向'],
        ['3428 draft.docx', '用能力等级和评估维度衡量模块成熟度'],
        ['Tech-discussions.docx', '对齐 FDE、OKF、Knowledge Repository、Prompt Studio、KSTAR 和 FDE Workbench 路线'],
    ]
)

# ---- Save ----
output_path = r'C:\Users\尹镇宇\Desktop\Elite20\nseap-knowledge-base\PRODUCT-SPEC.docx'
doc.save(output_path)
print(f'OK: {output_path}')
print(f'File size: {os.path.getsize(output_path)} bytes')
